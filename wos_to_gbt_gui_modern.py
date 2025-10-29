#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
wos_to_gbt_gui_modern.py

目标：与 Python 3.9 + 新版 pandas 配合良好，尽量使用现代库读取 Excel/CSV。
策略：
 - 对 .xlsx 使用 openpyxl（pandas 默认通常可以自动使用）
 - 对 .xls 尝试使用 xlrd（若用户安装的是 xlrd>=2.0，则说明不支持 .xls；脚本会尝试用 pyexcel-xls 回退）
 - 对 .xlsb 尝试使用 pyxlsb（如果安装）
 - 对 .csv 使用 pandas.read_csv（尝试自动检测编码）
 - 在无法读取 .xls 时给出清晰升级/降级建议（如何安装兼容的 xlrd 或将文件另存为 .xlsx）

使用方式：双击运行 → 选择文件（xlsx/xls/xlsb/csv）→ 在同目录生成：
 - 参考文献输出.docx
 - 参考文献输出.xlsx

依赖（请先安装）:
    pip install pandas openpyxl python-docx chardet pyxlsb pyexcel-xls

注：如果你仍然遇到 .xls 读取问题，最稳妥的方法是：
    1) 在 Excel 打开该文件，另存为 .xlsx
    2) 或者安装旧版 xlrd：pip install xlrd==1.2.0
"""
import re, sys, traceback, zipfile
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path
import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

# 小型中文姓氏集合（拼音）
CHINESE_SURNAME_SET = {
    'li','wang','zhang','liu','chen','yang','zhao','huang','zhou','wu','xu','sun','ma','zhu','hu',
    'guo','he','gao','lin','luo','zheng','liang','xie','song','tang','feng','yu','dong','kou','cao',
    'pan','wei','jiang','han','xiao','du','ye','pei'
}

def is_cjk(s: str) -> bool:
    if not isinstance(s, str):
        return False
    return any('\u4e00' <= ch <= '\u9fff' for ch in s)

def normalize_spaces(s: str):
    return re.sub(r'\s+', ' ', s.strip()) if isinstance(s, str) else ''

def format_name_smart(name: str):
    if not isinstance(name, str):
        return ''
    s = normalize_spaces(name.replace('.', ''))
    if not s:
        return ''
    if is_cjk(s):
        return s
    if ',' in s:
        parts = [p.strip() for p in s.split(',') if p.strip()]
        surname = parts[0]
        rest = ' '.join(parts[1:]) if len(parts) > 1 else ''
        initials = ''.join([tok[0].upper() for tok in rest.split() if tok])
        return f"{surname} {' '.join(list(initials))}".strip() if initials else surname
    toks = s.split()
    first_lower = toks[0].lower()
    if first_lower in CHINESE_SURNAME_SET:
        surname = toks[0]
        initials = ''.join([tok[0].upper() for tok in toks[1:] if tok])
        return f"{surname} {' '.join(list(initials))}".strip() if initials else surname
    if len(toks) == 2:
        if len(toks[1]) == 1:
            surname = toks[0]
            initials = toks[1].upper()
            return f"{surname} {initials}"
        if len(toks[0]) <= 3:
            surname = toks[0]
            initials = toks[1][0].upper()
            return f"{surname} {initials}"
        surname = toks[-1]
        initials = toks[0][0].upper()
        return f"{surname} {initials}"
    surname = toks[-1]
    initials = ''.join([tok[0].upper() for tok in toks[:-1] if tok])
    return f"{surname} {' '.join(list(initials))}".strip() if initials else surname

def split_authors_raw(authors_raw: str):
    if not isinstance(authors_raw, str):
        return []
    s = authors_raw.strip()
    if not s:
        return []
    if ';' in s:
        parts = [p.strip() for p in s.split(';') if p.strip()]
        return parts
    if '\n' in s:
        parts = [p.strip() for p in s.splitlines() if p.strip()]
        return parts
    if '；' in s:
        parts = [p.strip() for p in s.split('；') if p.strip()]
        return parts
    if re.search(r'\band\b', s, re.IGNORECASE):
        parts = [p.strip() for p in re.split(r'\band\b', s, flags=re.IGNORECASE) if p.strip()]
        return parts
    if ' & ' in s:
        parts = [p.strip() for p in s.split(' & ') if p.strip()]
        return parts
    if s.count(',') >= 2 and ';' not in s:
        parts = [p.strip() for p in re.split(r',\s*', s) if p.strip()]
        return parts
    parts = [p.strip() for p in re.split(r'[,/]', s) if p.strip()]
    return parts

def format_authors_block(authors_raw: str, truncate_n: int = 3):
    authors_list = split_authors_raw(authors_raw)
    if not authors_list:
        return '', False
    any_cjk = any(is_cjk(a) for a in authors_list)
    formatted = []
    for a in authors_list:
        a = a.strip()
        if not a:
            continue
        if is_cjk(a):
            formatted.append(a)
        else:
            formatted.append(format_name_smart(a))
    if len(formatted) > truncate_n:
        first = formatted[:truncate_n]
        if any_cjk:
            return ','.join(first) + ',等.', True
        else:
            return ', '.join(first) + ', et al.', False
    if any_cjk:
        return ','.join(formatted), True
    else:
        return ', '.join(formatted), False

def safe_get(row, variants):
    for v in variants:
        if v in row and pd.notna(row[v]):
            val = row[v]
            if isinstance(val, float) and val.is_integer():
                return str(int(val))
            return str(val).strip()
    return ''

def build_volpage_str(volume, issue, artnum, pages):
    volissue = ''
    if volume and issue:
        volissue = f"{volume}({issue})"
    elif volume:
        volissue = f"{volume}"
    elif issue:
        volissue = f"({issue})"
    pagepart = ''
    candidate = artnum or pages
    if candidate:
        candidate = str(candidate).strip()
        candidate = re.sub(r'pp\.\s*', '', candidate, flags=re.IGNORECASE)
        pagepart = f":{candidate}"
    if volissue and pagepart:
        return f"{volissue}{pagepart}"
    elif volissue:
        return volissue
    elif pagepart:
        return pagepart.lstrip(':')
    else:
        return ''

def format_reference_from_row(row, idx, truncate_n=3):
    pubtype = safe_get(row, ['Publication Type','publication type','PubType','Record Type','Document Type'])
    authors_raw = safe_get(row, ['Authors','Author','AU','authors','Author(s)'])
    title = safe_get(row, ['Article Title','Title','article title','TI'])
    source = safe_get(row, ['Source Title','Journal','Source','Source title','SO'])
    volume = safe_get(row, ['Volume','volume','VL'])
    issue = safe_get(row, ['Issue','issue','IS'])
    artnum = safe_get(row, ['Article Number','Art. No.','art. no.','Article Number','ArticleNumber'])
    pages = safe_get(row, ['Pages','Page','PG'])
    doi = safe_get(row, ['DOI','doi','DOI:'])
    year = safe_get(row, ['Publication Year','Year','PY','publication year'])

    authors_block, is_chinese_chars = format_authors_block(authors_raw, truncate_n=truncate_n)
    pubcode = '[J]'
    volpage = build_volpage_str(volume, issue, artnum, pages)

    prefix = f"[{idx}] "
    if is_chinese_chars:
        parts = f"{prefix}{authors_block}. {title}{pubcode} {source}, {year}"
        if volpage:
            parts += f", {volpage}"
        if doi:
            parts += f", doi:{doi}"
        if not parts.endswith("."):
            parts += "."
        return parts
    else:
        volpage_str = f", {volpage}" if volpage else ""
        doi_str = f". doi:{doi}" if doi else ""
        parts = f"{prefix}{authors_block}. {title}{pubcode} {source}, {year}{volpage_str}{doi_str}."
        parts = re.sub(r'\s+,', ',', parts)
        parts = re.sub(r'\.\s*\.', '.', parts)
        parts = re.sub(r'\s{2,}', ' ', parts)
        return parts

def read_excel_modern(path: Path):
    suffix = path.suffix.lower()
    # CSV
    if suffix == '.csv':
        try:
            import chardet
            raw = path.read_bytes()
            enc = chardet.detect(raw)['encoding'] or 'utf-8'
        except Exception:
            enc = 'utf-8'
        return pd.read_csv(path, encoding=enc)
    # xlsx
    if suffix == '.xlsx' or suffix == '.xlsm' or suffix == '.ods':
        return pd.read_excel(path, engine='openpyxl')
    # xlsb
    if suffix == '.xlsb':
        try:
            return pd.read_excel(path, engine='pyxlsb')
        except Exception as e:
            raise RuntimeError("Failed to read .xlsb file. Please install pyxlsb (pip install pyxlsb)")
    # xls (legacy BIFF)
    if suffix == '.xls':
        # try pandas default first (may require xlrd)
        try:
            return pd.read_excel(path, engine='xlrd')
        except Exception as e_xlrd:
            # try pyexcel-xls as fallback (if installed)
            try:
                from pyexcel_xls import get_data as px_get_data
                data = px_get_data(str(path))
                # take first sheet
                first_sheet = next(iter(data))
                rows = data[first_sheet]
                df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame()
                return df
            except Exception:
                # give user clear guidance
                raise RuntimeError(
                    "Failed to read .xls with xlrd. Modern xlrd (>=2.0) dropped .xls support.\n"
                    "Options:\n"
                    " 1) Install a legacy xlrd that supports .xls: pip install xlrd==1.2.0\n"
                    " 2) Convert the .xls to .xlsx in Excel and re-run\n"
                    " 3) Install pyexcel-xls: pip install pyexcel-xls and try again (script attempted fallback but it failed)."
                )
    # unknown suffix: try pandas to infer
    try:
        return pd.read_excel(path)
    except Exception as e:
        # fallback to csv attempt
        try:
            import chardet
            raw = path.read_bytes()
            enc = chardet.detect(raw)['encoding'] or 'utf-8'
            return pd.read_csv(path, encoding=enc)
        except Exception:
            raise RuntimeError(f"Unable to read file {path}. Error: {e}")

def generate_outputs(input_path: Path, truncate_n: int = 3):
    df = read_excel_modern(input_path)
    # Basic validation
    existing = set(df.columns.astype(str))
    expected_cols = {'Authors','Author','AU','Article Title','Title','Source Title','Journal','Publication Year','Year'}
    if not any(col in existing for col in expected_cols):
        sample_cols = ', '.join(list(existing)[:10])
        sample_path = input_path.parent / 'wos_input_sample.csv'
        try:
            df.head(10).to_csv(sample_path, index=False, encoding='utf-8-sig')
        except Exception:
            sample_path = None
        raise RuntimeError(f"Uploaded file doesn't contain expected columns. Found (first 10): {sample_cols}. Saved sample to: {sample_path}" if sample_path else "")
    citations = []
    for i, row in df.iterrows():
        citations.append(format_reference_from_row(row, i+1, truncate_n=truncate_n))
    out_dir = input_path.parent
    out_docx = out_dir / '参考文献输出.docx'
    out_xlsx = out_dir / '参考文献输出.xlsx'
    pd.DataFrame({'FormattedReference': citations}).to_excel(out_xlsx, index=False)
    if Document is None:
        raise RuntimeError("python-docx not installed. Please install: pip install python-docx")
    doc = Document()
    style = doc.styles['Normal']
    try:
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
    except Exception:
        pass
    doc.add_heading('参考文献', level=1)
    for c in citations:
        p = doc.add_paragraph(c)
    doc.save(out_docx)
    return out_docx, out_xlsx

def run_gui():
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo("选择文件", "请选择你的 WoS 导出的文件（.xlsx/.xls/.xlsb/.csv）。")
    file_path = filedialog.askopenfilename(title="选择你的WOS导出文件", filetypes=[("Excel/CSV", "*.xlsx;*.xls;*.xlsb;*.csv"), ("All Files","*.*")])
    if not file_path:
        messagebox.showinfo("取消", "未选择文件，已退出。")
        return
    input_path = Path(file_path)
    try:
        out_docx, out_xlsx = generate_outputs(input_path, truncate_n=3)
        messagebox.showinfo("完成", f"已生成文件：\n{out_docx}\n{out_xlsx}")
    except Exception as e:
        # show full error to user and give suggestions
        msg = f"生成失败：\n{e}\n\n建议：\n - 若为 .xls 文件，请在 Excel 中打开并另存为 .xlsx，或安装兼容 xlrd: pip install xlrd==1.2.0\n - 若为 .xlsb 文件，请安装 pyxlsb: pip install pyxlsb\n - 若为 .csv 文件，请确保编码为 UTF-8 或 UTF-8-SIG\n\n如果你愿意，把该文件发给我（或把错误信息粘贴过来），我可以为你直接处理。"
        if isinstance(e, RuntimeError):
            messagebox.showerror("错误", msg)
        else:
            messagebox.showerror("错误", f"生成失败：\n{e}")
if __name__ == "__main__":
    run_gui()